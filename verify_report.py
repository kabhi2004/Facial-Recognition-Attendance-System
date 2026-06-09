import docx

def verify():
    filename = "Facial_Recognition_Attendance_System_Report.docx"
    try:
        doc = docx.Document(filename)
    except Exception as e:
        print(f"Error opening report: {e}")
        return

    paragraphs = doc.paragraphs
    tables = doc.tables
    
    total_words = 0
    for p in paragraphs:
        total_words += len(p.text.split())
        
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total_words += len(p.text.split())
                    
    print("\n--- REPORT AUDIT STATISTICS ---")
    print(f"Report Filename:   {filename}")
    print(f"Total Paragraphs:  {len(paragraphs)}")
    print(f"Total Tables:      {len(tables)}")
    print(f"Estimated Words:   {total_words}")
    
    # 1 page of body text in double-spacing / 1.5-spacing with margins is ~350 words.
    # Tables also occupy significant page space.
    est_pages = (total_words / 320) + (len(tables) * 0.5) + 6 # 6 pages of front matter & breaks
    print(f"Estimated Pages:   {est_pages:.1f} pages")
    print("--------------------------------\n")

if __name__ == "__main__":
    verify()
